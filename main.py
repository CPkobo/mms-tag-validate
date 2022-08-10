import csv
from docx import Document
import os
import re


def validate_docx_file(path, tag, index_col=1, s_col=2, t_col=3):
    document = Document(path)
    result_ = ['-----', path, '-----']
    # First three tables are for memsource system
    table_count = len(document.tables) - 3
    for table_index, table in enumerate(document.tables[3:]):
        row_count = len(table.rows)
        table_index_ = table_index + 1
        for row_index, row in enumerate(table.rows):
            if row_index % 100 == 0:
                print(f'{row_index} / {row_count} @Table {table_index_} / {table_count}')
            cells = row.cells
            index = cells[index_col].text
            src_text = cells[s_col].text
            tgt_text = cells[t_col].text
            if tgt_text == '':
                continue

            row_result = [index]
            src_tags = tag.findall(src_text)
            tgt_tags = tag.findall(tgt_text)

            if len(src_tags) != len(tgt_tags):
                row_result.append(
                    'Tag Nums Unmatch:'
                    + str(len(src_tags))
                    + ' vs '
                    + str(len(tgt_tags))
                )

            err_tags = []
            tgt_tag_str = '|'.join(tgt_tags)

            for t in src_tags:
                if t in tgt_tag_str:
                    tgt_tag_str = tgt_tag_str.replace(t, '', 1)
                else:
                    err_tags.append(t)

            if len(err_tags) > 0:
                err_tags_str = ','.join(err_tags)
                row_result.append(f'Only in Source: {err_tags_str}')
            remained = tgt_tag_str.replace('|', '')
            if remained != '':
                tgt_remained = [x for x in tgt_tag_str.split('|') if x != '']
                tgt_remained_str = ','.join(tgt_remained)
                row_result.append(f'Only in Target: {tgt_remained_str}')

            if len(row_result) > 1:
                result_.append('\t'.join(row_result))

    return result_


if __name__ == '__main__':
    f = open('./reg.txt', 'r', encoding='utf-8')
    tag_exp = re.compile(f.read().strip())
    results = []
    local_dir = './_local'
    files = os.listdir(local_dir)
    for file in files:
        if (file.endswith('.docx')) and (file.startswith('$') is False):
            result = validate_docx_file(f'{local_dir}/{file}', tag_exp)
            results.append(result)
    with open('_local/result.tsv', 'w', encoding='utf-8') as f:
        for res in results:
            f.write('\n'.join(res))
            f.write('\n')

